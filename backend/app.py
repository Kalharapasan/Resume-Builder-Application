from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io
import os
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import spacy
from werkzeug.utils import secure_filename
import json

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'doc', 'docx'}
MAX_FILE_SIZE = 16 * 1024 * 1024

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
    print("✓ spaCy model loaded successfully")
except Exception as e:
    print(f"⚠ Warning: spaCy model not loaded: {e}")
    nlp = None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


class EnhancedResumeParser:
    """Enhanced resume parser with improved accuracy"""
    
    def __init__(self, text):
        self.text = text
        self.lines = [line.strip() for line in text.split('\n') if line.strip()]
        
    def extract_email(self):
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, self.text)
        return matches[0] if matches else ""
    
    def extract_phone(self):
        phone_patterns = [
            r'\+?1?\s*\(?(\d{3})\)?[\s.-]?(\d{3})[\s.-]?(\d{4})',
            r'\+?\d{1,3}[\s.-]?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{4}',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, self.text)
            if match:
                return match.group(0)
        return ""
    
    def extract_name(self):
        if self.lines:
            first_line = self.lines[0]
            first_line = re.sub(r'^(Resume|CV|Curriculum Vitae)[\s:]*', '', first_line, flags=re.IGNORECASE)
            if len(first_line.split()) <= 5 and first_line[0].isupper():
                return first_line
        return "Your Name"
    
    def extract_linkedin(self):
        """Extract LinkedIn profile URL"""
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        match = re.search(linkedin_pattern, self.text, re.IGNORECASE)
        return match.group(0) if match else ""
    
    def extract_location(self):
        location_patterns = [
            r'\b[A-Z][a-z]+,\s*[A-Z]{2}\b',
            r'\b[A-Z][a-z]+,\s*[A-Z][a-z]+\b'
        ]
        for pattern in location_patterns:
            match = re.search(pattern, self.text)
            if match:
                return match.group(0)
        return ""
    
    def extract_skills(self):
        skills = []
        tech_skills = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'React', 'Angular', 'Vue',
            'Node.js', 'Express', 'Django', 'Flask', 'FastAPI', 'Spring Boot',
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Elasticsearch',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'Git',
            'HTML', 'CSS', 'SCSS', 'Tailwind', 'Bootstrap', 'Material-UI',
            'REST API', 'GraphQL', 'gRPC', 'WebSocket',
            'Machine Learning', 'Deep Learning', 'AI', 'Data Science', 
            'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy',
            'C++', 'C#', 'Ruby', 'PHP', 'Swift', 'Kotlin', 'Go', 'Rust',
            'Agile', 'Scrum', 'DevOps', 'CI/CD', 'Microservices', 'Linux',
            'Firebase', 'Supabase', 'Next.js', 'Svelte', 'Remix'
        ]
        
        text_lower = self.text.lower()
        skills_keywords = ['skills', 'technical skills', 'competencies', 'expertise', 'technologies', 'tools']
        
        for keyword in skills_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                skills_section = self.text[idx:idx+600]
                for skill in tech_skills:
                    if re.search(r'\b' + re.escape(skill) + r'\b', skills_section, re.IGNORECASE):
                        if skill not in skills:
                            skills.append(skill)
                break
        
        if not skills:
            for skill in tech_skills[:20]:
                if re.search(r'\b' + re.escape(skill) + r'\b', self.text, re.IGNORECASE):
                    if skill not in skills:
                        skills.append(skill)
        
        skills = list(dict.fromkeys(skills))[:15]
        return skills if skills else ["Communication", "Problem Solving", "Teamwork", "Leadership"]
    
    def extract_experience(self):
        experience = []
        text_lower = self.text.lower()
        exp_keywords = ['experience', 'work experience', 'employment', 'work history', 'professional experience']
        
        exp_start_idx = -1
        for keyword in exp_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                exp_start_idx = idx
                break
        
        if exp_start_idx != -1:
            end_keywords = ['education', 'skills', 'projects', 'certifications']
            exp_end_idx = len(self.text)
            
            for keyword in end_keywords:
                idx = text_lower.find(keyword, exp_start_idx + 50)
                if idx != -1 and idx < exp_end_idx:
                    exp_end_idx = idx
            
            exp_section = self.text[exp_start_idx:exp_end_idx]
            date_pattern = r'(\d{4}|[A-Z][a-z]+\s+\d{4})\s*[-–—]\s*(\d{4}|Present|Current|Now)'
            dates = re.findall(date_pattern, exp_section, re.IGNORECASE)
            
            if dates:
                for i, date in enumerate(dates[:4]):
                    lines_around_date = exp_section.split('\n')
                    title = f"Position {i+1}"
                    company = f"Company {i+1}"
                    
                    for line in lines_around_date:
                        if date[0] in line:
                            idx = lines_around_date.index(line)
                            if idx > 0:
                                title = lines_around_date[idx-1].strip()
                            if idx > 1:
                                company = lines_around_date[idx-2].strip()
                            break
                    
                    experience.append({
                        'title': title[:100],
                        'company': company[:100],
                        'period': f"{date[0]} - {date[1]}",
                        'description': 'Responsible for key tasks and deliverables in the role'
                    })
        
        if not experience:
            experience = [
                {
                    'title': 'Software Developer',
                    'company': 'Tech Company',
                    'period': '2020 - Present',
                    'description': 'Developed and maintained web applications using modern technologies'
                }
            ]
        
        return experience[:4]
    
    def extract_education(self):
        education = []
        text_lower = self.text.lower()
        edu_keywords = ['education', 'academic', 'qualification', 'degree']
        
        edu_start_idx = -1
        for keyword in edu_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                edu_start_idx = idx
                break
        
        if edu_start_idx != -1:
            edu_section = self.text[edu_start_idx:edu_start_idx+600]
            
            degree_patterns = [
                r'(Bachelor|Master|PhD|Doctorate|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|B\.?Tech|M\.?Tech|MBA|BBA|BCA|MCA)[\s\w\.,\(\)]*',
                r'(Diploma|Certificate|Associate)[\s\w]*'
            ]
            
            degrees = []
            for pattern in degree_patterns:
                matches = re.findall(pattern, edu_section, re.IGNORECASE)
                degrees.extend(matches)
            
            year_pattern = r'\b(19|20)\d{2}\b'
            years = re.findall(year_pattern, edu_section)
            
            institution_pattern = r'(?:at |from |,\s*)([A-Z][A-Za-z\s&,\.]+(?:University|College|Institute|School|Academy))'
            institutions = re.findall(institution_pattern, edu_section, re.IGNORECASE)
            
            for i in range(min(len(degrees), 3)):
                education.append({
                    'degree': degrees[i].strip(),
                    'institution': institutions[i].strip() if i < len(institutions) else 'University',
                    'year': years[i] if i < len(years) else '2020'
                })
        
        if not education:
            education = [{
                'degree': 'Bachelor of Science in Computer Science',
                'institution': 'University',
                'year': '2020'
            }]
        
        return education
    
    def extract_summary(self):
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'professional summary', 'overview']
        text_lower = self.text.lower()
        
        for keyword in summary_keywords:
            idx = text_lower.find(keyword)
            if idx != -1:
                summary_start = idx + len(keyword)
                summary_section = self.text[summary_start:summary_start+500]
                sentences = re.split(r'[.!?]\s+', summary_section)
                valid_sentences = [s.strip() for s in sentences if len(s.split()) > 5]
                if valid_sentences:
                    summary = '. '.join(valid_sentences[:3]) + '.'
                    return summary[:400]
        
        return 'Experienced professional with strong technical skills and a proven track record of delivering high-quality results.'
    
    def calculate_score(self, data):
        """Calculate resume completeness score"""
        score = 0
        max_score = 100
        
        if data.get('name') and data['name'] != 'Your Name':
            score += 10
        if data.get('email'):
            score += 15
        if data.get('phone'):
            score += 10
        if data.get('summary') and len(data['summary']) > 50:
            score += 15
        if data.get('skills') and len(data['skills']) >= 5:
            score += 20
        if data.get('experience') and len(data['experience']) >= 1:
            score += 20
        if data.get('education') and len(data['education']) >= 1:
            score += 10
        
        return min(score, max_score)
    
    def parse(self):
        data = {
            'name': self.extract_name(),
            'email': self.extract_email(),
            'phone': self.extract_phone(),
            'linkedin': self.extract_linkedin(),
            'location': self.extract_location(),
            'summary': self.extract_summary(),
            'skills': self.extract_skills(),
            'experience': self.extract_experience(),
            'education': self.extract_education()
        }
        
        data['score'] = self.calculate_score(data)
        return data


class EnhancedPDFGenerator:
    """Enhanced PDF generator with new templates"""
    
    @staticmethod
    def create_modern_template(buffer, data):
        """Modern template with blue accent"""
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=30)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=28,
            textColor=colors.HexColor('#1e40af'), spaceAfter=8, fontName='Helvetica-Bold')
        
        contact_style = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=10,
            textColor=colors.HexColor('#4b5563'), spaceAfter=16)
        
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14,
            textColor=colors.HexColor('#1e40af'), spaceAfter=8, spaceBefore=14, fontName='Helvetica-Bold')
        
        story.append(Paragraph(data['name'], title_style))
        
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('location'):
            contact_parts.append(data['location'])
        
        if contact_parts:
            story.append(Paragraph(' | '.join(contact_parts), contact_style))
        
        story.append(Spacer(1, 0.1*inch))
        
        if data.get('summary'):
            story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            story.append(Paragraph(data['summary'], styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        if data.get('experience'):
            story.append(Paragraph('WORK EXPERIENCE', heading_style))
            for exp in data['experience']:
                job_style = ParagraphStyle('Job', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11)
                story.append(Paragraph(exp['title'], job_style))
                company_style = ParagraphStyle('Company', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#4b5563'))
                story.append(Paragraph(f"{exp['company']} | {exp['period']}", company_style))
                story.append(Paragraph(exp['description'], styles['Normal']))
                story.append(Spacer(1, 0.12*inch))
        
        if data.get('education'):
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in data['education']:
                degree_style = ParagraphStyle('Degree', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11)
                story.append(Paragraph(edu['degree'], degree_style))
                institution_style = ParagraphStyle('Institution', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#4b5563'))
                story.append(Paragraph(f"{edu['institution']} | {edu['year']}", institution_style))
                story.append(Spacer(1, 0.12*inch))
        
        if data.get('skills'):
            story.append(Paragraph('SKILLS', heading_style))
            skills_text = ' • '.join(data['skills'])
            story.append(Paragraph(skills_text, styles['Normal']))
        
        doc.build(story)
        return buffer
    
    @staticmethod
    def create_professional_template(buffer, data):
        """Professional centered template"""
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=60, leftMargin=60, topMargin=60, bottomMargin=30)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=26,
            textColor=colors.black, spaceAfter=8, alignment=TA_CENTER, fontName='Helvetica-Bold')
        
        contact_style = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=10,
            textColor=colors.HexColor('#374151'), spaceAfter=20, alignment=TA_CENTER)
        
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=12,
            textColor=colors.black, spaceAfter=6, spaceBefore=12, fontName='Helvetica-Bold')
        
        story.append(Paragraph(data['name'], title_style))
        
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('location'):
            contact_parts.append(data['location'])
        
        if contact_parts:
            story.append(Paragraph(' | '.join(contact_parts), contact_style))
        
        story.append(Spacer(1, 0.1*inch))
        
        if data.get('summary'):
            story.append(Paragraph('SUMMARY', heading_style))
            story.append(Paragraph(data['summary'], styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        if data.get('experience'):
            story.append(Paragraph('EXPERIENCE', heading_style))
            for exp in data['experience']:
                job_style = ParagraphStyle('Job', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11)
                story.append(Paragraph(exp['title'], job_style))
                company_style = ParagraphStyle('Company', parent=styles['Normal'], fontSize=10,
                    textColor=colors.HexColor('#4b5563'), fontName='Helvetica-Oblique')
                story.append(Paragraph(f"{exp['company']} | {exp['period']}", company_style))
                story.append(Paragraph(exp['description'], styles['Normal']))
                story.append(Spacer(1, 0.12*inch))
        
        if data.get('education'):
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in data['education']:
                degree_style = ParagraphStyle('Degree', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11)
                story.append(Paragraph(edu['degree'], degree_style))
                institution_style = ParagraphStyle('Institution', parent=styles['Normal'], fontSize=10, textColor=colors.HexColor('#4b5563'))
                story.append(Paragraph(f"{edu['institution']} | {edu['year']}", institution_style))
                story.append(Spacer(1, 0.12*inch))
        
        if data.get('skills'):
            story.append(Paragraph('SKILLS', heading_style))
            skills_text = ' • '.join(data['skills'])
            story.append(Paragraph(skills_text, styles['Normal']))
        
        doc.build(story)
        return buffer
    
    @staticmethod
    def create_creative_template(buffer, data):
        """Creative template with colorful design"""
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=30)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle('CreativeTitle', parent=styles['Heading1'], fontSize=30,
            textColor=colors.HexColor('#ec4899'), spaceAfter=8, fontName='Helvetica-Bold')
        
        heading_style = ParagraphStyle('CreativeHeading', parent=styles['Heading2'], fontSize=14,
            textColor=colors.HexColor('#f97316'), spaceAfter=8, spaceBefore=14, fontName='Helvetica-Bold')
        
        story.append(Paragraph(data['name'], title_style))
        
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        
        if contact_parts:
            story.append(Paragraph(' | '.join(contact_parts), styles['Normal']))
        
        story.append(Spacer(1, 0.1*inch))
        
        if data.get('summary'):
            story.append(Paragraph('ABOUT ME', heading_style))
            story.append(Paragraph(data['summary'], styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        if data.get('skills'):
            story.append(Paragraph('EXPERTISE', heading_style))
            skills_text = ' • '.join(data['skills'])
            story.append(Paragraph(skills_text, styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        if data.get('experience'):
            story.append(Paragraph('EXPERIENCE', heading_style))
            for exp in data['experience']:
                story.append(Paragraph(f"<b>{exp['title']}</b>", styles['Normal']))
                story.append(Paragraph(f"{exp['company']} | {exp['period']}", styles['Normal']))
                story.append(Paragraph(exp['description'], styles['Normal']))
                story.append(Spacer(1, 0.12*inch))
        
        if data.get('education'):
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in data['education']:
                story.append(Paragraph(f"<b>{edu['degree']}</b>", styles['Normal']))
                story.append(Paragraph(f"{edu['institution']} | {edu['year']}", styles['Normal']))
                story.append(Spacer(1, 0.12*inch))
        
        doc.build(story)
        return buffer


class DocxGenerator:
    """Generate DOCX resumes"""
    
    @staticmethod
    def create_docx(data):
        doc = Document()
        
        # Title
        title = doc.add_heading(data['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('location'):
            contact_parts.append(data['location'])
        
        if contact_parts:
            p = doc.add_paragraph(' | '.join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if data.get('summary'):
            doc.add_heading('Professional Summary', 1)
            doc.add_paragraph(data['summary'])
        
        # Skills
        if data.get('skills'):
            doc.add_heading('Skills', 1)
            doc.add_paragraph(' • '.join(data['skills']))
        
        # Experience
        if data.get('experience'):
            doc.add_heading('Work Experience', 1)
            for exp in data['experience']:
                doc.add_heading(exp['title'], 2)
                doc.add_paragraph(f"{exp['company']} | {exp['period']}")
                doc.add_paragraph(exp['description'])
        
        # Education
        if data.get('education'):
            doc.add_heading('Education', 1)
            for edu in data['education']:
                doc.add_heading(edu['degree'], 2)
                doc.add_paragraph(f"{edu['institution']} | {edu['year']}")
        
        return doc


@app.route('/api/upload', methods=['POST'])
def upload_document():
    """Handle document upload and parsing"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only .doc and .docx files are supported'}), 400
    
    try:
        doc = Document(file)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += '\n' + cell.text
        
        if not text.strip():
            return jsonify({'error': 'Document appears to be empty'}), 400
        
        parser = EnhancedResumeParser(text)
        parsed_data = parser.parse()
        
        return jsonify({
            'success': True,
            'data': parsed_data,
            'raw_text': text[:500]
        })
    
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return jsonify({'error': f'Error processing document: {str(e)}'}), 500


@app.route('/api/update-data', methods=['POST'])
def update_data():
    """Update parsed resume data"""
    data = request.json
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    return jsonify({'success': True, 'data': data})


@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    """Generate PDF resume"""
    data = request.json
    resume_data = data.get('resumeData')
    template = data.get('template', 'modern')
    
    if not resume_data:
        return jsonify({'error': 'No resume data provided'}), 400
    
    try:
        buffer = io.BytesIO()
        
        if template == 'professional':
            EnhancedPDFGenerator.create_professional_template(buffer, resume_data)
        elif template == 'creative':
            EnhancedPDFGenerator.create_creative_template(buffer, resume_data)
        else:
            EnhancedPDFGenerator.create_modern_template(buffer, resume_data)
        
        buffer.seek(0)
        
        name = resume_data.get('name', 'Resume').replace(' ', '_')
        filename = f"{name}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)
    
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return jsonify({'error': f'Error generating PDF: {str(e)}'}), 500


@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    """Generate DOCX resume"""
    data = request.json
    resume_data = data.get('resumeData')
    
    if not resume_data:
        return jsonify({'error': 'No resume data provided'}), 400
    
    try:
        doc = DocxGenerator.create_docx(resume_data)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        name = resume_data.get('name', 'Resume').replace(' ', '_')
        filename = f"{name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True, download_name=filename)
    
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return jsonify({'error': f'Error generating DOCX: {str(e)}'}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'version': '2.0',
        'message': 'Resume Generator API is running',
        'spacy_loaded': nlp is not None,
        'timestamp': datetime.now().isoformat()
    })


if __name__ == '__main__':
    print("\n" + "="*60)
    print("🚀 AI Resume Generator v2.0 - API Starting...")
    print("="*60)
    print(f"📁 Upload folder: {UPLOAD_FOLDER}")
    print(f"🤖 spaCy loaded: {'✓ Yes' if nlp else '✗ No'}")
    print(f"🌐 Server: http://localhost:5000")
    print(f"💚 Health: http://localhost:5000/api/health")
    print(f"📊 Version: 2.0 (Enhanced)")
    print("="*60 + "\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)